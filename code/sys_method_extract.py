import os
import re
import json
import time
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from google import genai
from google.genai import types
from openpyxl.styles import Alignment

# Configuration
API_KEY = os.getenv("GEMINI_API_KEY", "YOUR_API_KEY")
MODEL_NAME = "gemini-2.5-flash"

BASE_DIR = r"YOUR_IN_PATH"
OUTPUT_FILE = r"YOUR_OUT_PATH"

MAX_PDF_PAGES_TO_SCAN = 80
MAX_BLOCKS_TO_SEND = 18
MAX_CHARS_PER_BLOCK = 15000
SLEEP_SECONDS = 1.0

DEBUG_DIR = Path("debug_gemini")
DEBUG_DIR.mkdir(exist_ok=True)

client = genai.Client(api_key=API_KEY)

ILLEGAL_CHARACTERS_RE = re.compile(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]")

SECTION_KEYWORDS = [
    "experimental",
    "experiment",
    "materials and methods",
    "methods",
    "synthesis",
    "synthetic procedure",
    "preparation",
    "catalyst preparation",
    "sample preparation",
    "material synthesis",
    "fabrication",
    "reagents",
    "chemicals",
]

EXCLUDE_HINTS = [
    "electrochemical measurements",
    "electrochemical test",
    "orr",
    "rde",
    "lsv",
    "cvs",
    "zab",
    "zinc-air",
    "characterization",
    "xrd",
    "xps",
    "tem",
    "sem",
    "exafs",
    "xanes",
    "bet",
]


# Util
def clean_excel_string(x):
    if x is None:
        return ""
    if not isinstance(x, str):
        x = str(x)
    x = ILLEGAL_CHARACTERS_RE.sub("", x)
    return x.strip()


def safe_slug(name: str) -> str:
    name = re.sub(r"[^\w\-.]+", "_", name, flags=re.UNICODE)
    return name[:150].strip("_") or "debug"


def score_text_block(text: str) -> int:
    t = (text or "").lower()
    score = 0

    for kw in SECTION_KEYWORDS:
        if kw in t:
            score += 3

    for kw in EXCLUDE_HINTS:
        if kw in t:
            score -= 1

    chemistry_patterns = [
        r"\bmg\b", r"\bg\b", r"\bml\b", r"\bmL\b", r"\bmmol\b",
        r"\b°C\b", r"\bK\b", r"\bh\b", r"\bmin\b",
        r"\bstir", r"\bsonicat", r"\bwash", r"\bdry", r"\bcalc", r"\bpyro"
    ]
    for pat in chemistry_patterns:
        if re.search(pat, t, flags=re.I):
            score += 1

    return score


def sanitize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    for col in df.columns:
        df[col] = df[col].map(clean_excel_string)
    return df


def parse_gemini_response(response, debug_name="debug_response"):
    parsed = getattr(response, "parsed", None)
    if parsed is not None:
        if isinstance(parsed, dict):
            return parsed
        try:
            return json.loads(json.dumps(parsed, ensure_ascii=False))
        except Exception:
            pass

    text = getattr(response, "text", None)
    if text:
        raw_text = text.strip()

        if raw_text.startswith("```"):
            raw_text = re.sub(r"^```(?:json)?\s*", "", raw_text)
            raw_text = re.sub(r"\s*```$", "", raw_text)

        try:
            return json.loads(raw_text)
        except Exception:
            pass

        m = re.search(r"\{.*\}", raw_text, flags=re.S)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass

    debug_base = safe_slug(debug_name)
    with open(DEBUG_DIR / f"{debug_base}.txt", "w", encoding="utf-8") as f:
        f.write(getattr(response, "text", "") or "[EMPTY response.text]")

    with open(DEBUG_DIR / f"{debug_base}_repr.txt", "w", encoding="utf-8") as f:
        f.write(repr(response))

    return None


# Read PDF / DOCX / DOC
def extract_pdf_text_blocks(file_path: Path, max_pages=MAX_PDF_PAGES_TO_SCAN, max_blocks=MAX_BLOCKS_TO_SEND):
    doc = fitz.open(file_path)
    blocks = []

    try:
        for i in range(min(len(doc), max_pages)):
            page = doc[i]
            text = page.get_text("text")
            text = clean_excel_string(text)

            if not text:
                continue

            blocks.append({
                "unit": f"PAGE {i + 1}",
                "text": text,
                "score": score_text_block(text),
                "page_index": i + 1,
            })
    finally:
        doc.close()

    if not blocks:
        return []

    hit_blocks = [b for b in blocks if b["score"] > 0]
    selected = hit_blocks if hit_blocks else blocks

    selected = sorted(selected, key=lambda x: (-x["score"], x["page_index"]))
    selected = sorted(selected[:max_blocks], key=lambda x: x["page_index"])
    return selected


def extract_docx_text_blocks(file_path: Path, max_blocks=MAX_BLOCKS_TO_SEND):
    doc = Document(file_path)
    blocks = []

    para_idx = 0
    for para in doc.paragraphs:
        text = clean_excel_string(para.text)
        if text:
            para_idx += 1
            blocks.append({
                "unit": f"PARA {para_idx}",
                "text": text,
                "score": score_text_block(text),
                "order": para_idx,
            })

    table_idx = 0
    for table in doc.tables:
        table_idx += 1
        row_texts = []
        for row in table.rows:
            cells = [clean_excel_string(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                row_texts.append(" | ".join(cells))
        if row_texts:
            text = "\n".join(row_texts)
            blocks.append({
                "unit": f"TABLE {table_idx}",
                "text": text,
                "score": score_text_block(text),
                "order": 100000 + table_idx,
            })

    if not blocks:
        return []

    hit_blocks = [b for b in blocks if b["score"] > 0]
    selected = hit_blocks if hit_blocks else blocks
    selected = sorted(selected, key=lambda x: (-x["score"], x.get("order", 0)))
    return selected[:max_blocks]


def convert_doc_to_docx(doc_path: Path) -> Path:
    try:
        import win32com.client
    except ImportError:
        raise RuntimeError("pywin32 not installed, failed to convert .doc file.")

    tmp_dir = Path(tempfile.mkdtemp(prefix="doc_convert_"))
    out_path = tmp_dir / f"{doc_path.stem}.docx"

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    try:
        doc = word.Documents.Open(str(doc_path.resolve()))
        doc.SaveAs(str(out_path), FileFormat=16)  # wdFormatXMLDocument
        doc.Close(False)
    finally:
        word.Quit()

    return out_path


def extract_text_blocks(file_path: Path):
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_text_blocks(file_path)

    if suffix == ".docx":
        return extract_docx_text_blocks(file_path)

    if suffix == ".doc":
        converted = convert_doc_to_docx(file_path)
        return extract_docx_text_blocks(converted)

    raise ValueError(f"Unsupported file type: {suffix}")


# Prompt (can be changed depend on task)
def build_prompt_from_blocks(blocks, source_name):
    text_parts = []
    for b in blocks:
        txt = (b["text"] or "")[:MAX_CHARS_PER_BLOCK]
        text_parts.append(f"\n===== {b['unit']} =====\n{txt}")

    joined_text = "\n".join(text_parts)

    system_prompt = (
        "You are a chemistry literature extraction engine. "
        "Your only job is to locate and copy original text from the source document. "
        "Do NOT summarize. Do NOT paraphrase. Do NOT rewrite. Do NOT shorten. "
        "Preserve the original wording as faithfully as possible."
    )

    user_prompt = f"""
Copy original synthesis-related text from the source document.

Source file:
{source_name}

Return JSON with:
- procedures: an array of objects
- each object must contain:
  - target_product
  - full_procedure
  - yield

Mandatory rules:
1. full_procedure must be copied from the source text only.
2. Do NOT summarize.
3. Do NOT paraphrase.
4. Do NOT rewrite.
5. Do NOT shorten.
6. Prefer a complete paragraph.
7. If one synthesis description spans multiple consecutive paragraphs, copy all consecutive relevant paragraphs together into full_procedure.
8. Keep synthesis details such as reagent names, amounts, solvents, temperatures, times, washing, drying, calcination, pyrolysis, centrifugation, stirring, sonication, precursor preparation, and support preparation.
9. Exclude characterization-only text, electrochemical testing, ORR/RDE/ZAB testing, device assembly, and performance discussion.
10. If a paragraph contains synthesis plus a small amount of nearby non-synthesis context, keep the original synthesis-containing paragraph rather than shortening it.
11. If no synthesis text is found for a product, do not invent one.
12. Return valid JSON only.

Paper content:
{joined_text}
"""
    return system_prompt, user_prompt


# LLM Extraction
def extract_synthesis_procedures(file_path: Path, client):
    print(f"Processing: {file_path.name}...", end=" ", flush=True)

    try:
        blocks = extract_text_blocks(file_path)
        if not blocks:
            print("No readable content.")
            return []
    except Exception as e:
        print(f"\n[Error] Reading file failed: {type(e).__name__}: {e}")
        return []

    system_prompt, user_prompt = build_prompt_from_blocks(blocks, file_path.name)

    response_schema = types.Schema(
        type=types.Type.OBJECT,
        properties={
            "procedures": types.Schema(
                type=types.Type.ARRAY,
                items=types.Schema(
                    type=types.Type.OBJECT,
                    properties={
                        "target_product": types.Schema(type=types.Type.STRING),
                        "full_procedure": types.Schema(type=types.Type.STRING),
                        "yield": types.Schema(type=types.Type.STRING),
                    },
                    required=["target_product", "full_procedure", "yield"],
                ),
            )
        },
        required=["procedures"],
    )

    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=user_prompt,
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                temperature=0.0,
                max_output_tokens=8192,
                response_mime_type="application/json",
                response_schema=response_schema,
            ),
        )

        result = parse_gemini_response(response, debug_name=file_path.stem)
        if not result:
            print("[Skip] Could not parse model response.")
            return []

        procedures = []
        for item in result.get("procedures", []):
            target_product = clean_excel_string(item.get("target_product", ""))
            full_procedure = clean_excel_string(item.get("full_procedure", ""))
            yield_text = clean_excel_string(item.get("yield", ""))

            if not full_procedure:
                continue

            procedures.append({
                "folder_path": str(file_path.parent),
                "source_file": file_path.name,
                "target_product": target_product,
                "yield": yield_text,
                "full_procedure": full_procedure,
            })

        if procedures:
            print(f"Done. ({len(procedures)} entries)")
        else:
            print("No procedures found.")

        return procedures

    except Exception as e:
        print(f"\n[Skip] API or JSON Error: {type(e).__name__}: {e}")
        return []


# Main
def main():
    if not API_KEY or API_KEY == "YOUR_API_KEY":
        raise ValueError("Please enter your GEMINI_API_KEY.")

    root_path = Path(BASE_DIR)
    output_path = Path(OUTPUT_FILE)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    supported_exts = {".pdf", ".docx", ".doc"}
    files = [p for p in root_path.rglob("*") if p.is_file() and p.suffix.lower() in supported_exts]

    if not files:
        print(f"No supported files found in {BASE_DIR}")
        return

    print(f"Starting batch processing of {len(files)} files...")

    all_data = []
    failed_files = []

    for file_path in files:
        data = extract_synthesis_procedures(file_path, client)
        if not data:
            failed_files.append(str(file_path))
        all_data.extend(data)
        time.sleep(SLEEP_SECONDS)

    if all_data:
        df = pd.DataFrame(all_data)

        cols = ["folder_path", "source_file", "target_product", "yield", "full_procedure"]
        df = df[[c for c in cols if c in df.columns]]

        df = sanitize_dataframe(df)
        df = df.drop_duplicates(
            subset=["source_file", "target_product", "full_procedure"],
            keep="first"
        ).reset_index(drop=True)

        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Sheet1")
            sheet = writer.sheets["Sheet1"]

            sheet.column_dimensions["A"].width = 45
            sheet.column_dimensions["B"].width = 45
            sheet.column_dimensions["C"].width = 25
            sheet.column_dimensions["D"].width = 18
            sheet.column_dimensions["E"].width = 120

            for row in sheet.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

        print(f"\nAll results saved to: {output_path}")
    else:
        print("\nNo data was extracted.")

    if failed_files:
        fail_txt = output_path.with_name("failed_files.txt")
        with open(fail_txt, "w", encoding="utf-8") as f:
            f.write("\n".join(failed_files))
        print(f"Failed/empty files logged to: {fail_txt}")

    print(f"Debug responses saved in: {DEBUG_DIR.resolve()}")


if __name__ == "__main__":
    main()
    